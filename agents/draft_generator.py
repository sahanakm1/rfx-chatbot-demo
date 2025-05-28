from agents.keyvault import get_secret
from azure.identity import DefaultAzureCredential
from azure.storage.blob import BlobServiceClient, ContainerClient, BlobBlock, BlobClient, StandardBlobTier
from agents.blob_storage import blob_storage

import os
import json
import io
from io import BytesIO
import uuid
import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

date_format="%Y_%m_%d_%H_%M_%S"
timestamp = datetime.datetime.now().strftime(date_format)

TOC = {
    "A. INTRODUCTION": ["A.1 JTI", "A.2 Our Engagement"],
    "B. PURPOSE OF THE RFP": ["B.1 Responses", "B.2 Schedule", "B.3 Queries", "B.4 Evaluation Criteria"],
    "C. CONTEXT": ["C.1 Project Scope and Objective", "C.2 JTI Requirements", "C.3 Proposal evaluation criteria"],
    "D. RESPONSE": ["D.1 Executive Summary", "D.2 Additional proposal details"],
    "E. APPENDICES": ["E.1 Supporting Documents"]
}

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p

def insert_toc(paragraph):
    """Insert a field code for TOC that Word will convert into a clickable TOC"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r = paragraph.add_run()
    r._r.append(fldChar1)
    r._r.append(instrText)
    r._r.append(fldChar2)
    r._r.append(fldChar3)

def build_doc_from_json(data_json, output_path="drafts/Generated_Document.docx"):
    doc = Document()
    doc.add_heading("Generated Proposal", level=0)
    
    # Insert TOC Placeholder
    toc_paragraph = doc.add_paragraph()
    insert_toc(toc_paragraph)
    doc.add_paragraph("")
    
    for section_title, subsections in TOC.items():
        section_key = section_title.split(".")[0]
        add_heading(doc, section_title, level=1)
        if subsections:
            for subsection in subsections:
                subsection_key = subsection.split()[0]
                heading_text = subsection
                content = data_json.get(section_key, {}).get(subsection_key, "(No content provided)")
                add_heading(doc, heading_text, level=2)
                add_paragraph(doc, content)
        else:
            content = data_json.get(section_key, "(No content provided)")
            add_paragraph(doc, content)
    
    # ✅ Create output directory if it doesn't exist
    #os.makedirs(os.path.dirname(output_path), exist_ok=True)   #Removed by Prerit
    
    
    # doc.save(output_path)      #Prerit
    # print(f"Document saved as {output_path}")  #Prerit
    # print("Open the document in Word and press F9 to update the TOC!")  #Prerit

    #output_path = os.path.abspath(output_path)      #Added & Removed by Prerit
    file_name = os.path.basename(output_path)        #Added by Prerit
    file_name = f"{file_name.split(".")[0]}_{timestamp}.{file_name.split('.')[1]}"
    #file_path = os.path.dirname(output_path)        #Added & Removed by Prerit

    word_buffer = BytesIO() #Added by Prerit
    doc.save(word_buffer) #Added by Prerit
    word_buffer.seek(0) #Added by Prerit
    
    blob_storage().writing_docx_file(container_name="rfx-draft", data=word_buffer,blob_name=file_name) #Added by Prerit

    #blob_storage().upload_blob_file(container_name="rfx-draft",file_path=file_path,file_name=file_name) #Added & Removed by Prerit
    
    return output_path