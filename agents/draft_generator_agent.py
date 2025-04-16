# agents/draft_generator_agent.py

from docx import Document
from datetime import datetime

def generate_rfx_draft(structured_brief: dict) -> str:
    """
    Generates an RFx Word document using the provided structured brief.
    Returns the path to the generated file.
    """
    doc = Document()
    doc.add_heading("RFx Document", level=0)

    # Add each section
    if structured_brief.get("background"):
        doc.add_heading("Background", level=1)
        doc.add_paragraph(structured_brief["background"])

    if structured_brief.get("objectives"):
        doc.add_heading("Objectives", level=1)
        doc.add_paragraph(structured_brief["objectives"])

    if structured_brief.get("scope"):
        doc.add_heading("Scope", level=1)
        doc.add_paragraph(structured_brief["scope"])

    if structured_brief.get("timeline"):
        doc.add_heading("Timeline", level=1)
        doc.add_paragraph(structured_brief["timeline"])

    if structured_brief.get("budget"):
        doc.add_heading("Budget", level=1)
        doc.add_paragraph(structured_brief["budget"])

    if structured_brief.get("evaluation_criteria"):
        doc.add_heading("Evaluation Criteria", level=1)
        doc.add_paragraph(structured_brief["evaluation_criteria"])

    if structured_brief.get("contact_info"):
        doc.add_heading("Contact Information", level=1)
        doc.add_paragraph(structured_brief["contact_info"])

    # Save the document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = f"rfx_draft_{timestamp}.docx"
    doc.save(output_path)
    return output_path